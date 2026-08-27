#!/usr/bin/env python3
"""Generate compassion-style group medical invoice docx (template clone)."""
from __future__ import annotations

import argparse
import copy
import json
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from lxml import etree

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
PATIENT_INFO_COLS = (0, 1, 2, 3, 4, 5)


def format_title_case_words(text) -> str:
    """First letter upper, remaining letters lower for each word."""
    raw = "" if text is None else str(text).strip()
    if not raw:
        return ""
    words = raw.split()
    formatted = []
    for word in words:
        if not word:
            continue
        if word.isdigit():
            formatted.append(word)
            continue
        if re.match(r"^\d", word):
            formatted.append(word)
            continue
        formatted.append(word[:1].upper() + word[1:].lower())
    return " ".join(formatted)


def format_sex(value) -> str:
    raw = "" if value is None else str(value).strip()
    if not raw:
        return ""
    return raw[:1].upper() + raw[1:].lower()


def format_age(value) -> str:
    raw = "" if value is None else str(value).strip()
    if not raw:
        return ""
    match = re.match(r"^(\d+)([A-Za-z]+)$", raw)
    if match:
        suffix = match.group(2)
        return f"{match.group(1)}{suffix[:1].upper()}{suffix[1:].lower()}"
    return format_title_case_words(raw)


def set_vmerge(tc, restart: bool) -> None:
    tc_pr = tc.find(f"{W}tcPr")
    if tc_pr is None:
        tc_pr = etree.Element(f"{W}tcPr")
        tc.insert(0, tc_pr)
    for old in tc_pr.findall(f"{W}vMerge"):
        tc_pr.remove(old)
    if restart:
        etree.SubElement(tc_pr, f"{W}vMerge", {f"{W}val": "restart"})
    else:
        etree.SubElement(tc_pr, f"{W}vMerge")


def merge_patient_info_columns(tbl) -> None:
    rows = tbl.findall(f"{W}tr")
    if len(rows) < 3:
        return
    data_rows = rows[1:]
    if len(data_rows) < 2:
        return
    for col_idx in PATIENT_INFO_COLS:
        for row_idx, tr in enumerate(data_rows):
            tcs = tr.findall(f"{W}tc")
            if col_idx >= len(tcs):
                continue
            set_vmerge(tcs[col_idx], restart=row_idx == 0)


def _ensure_tr_pr(tr):
    tr_pr = tr.find(f"{W}trPr")
    if tr_pr is None:
        tr_pr = etree.Element(f"{W}trPr")
        tr.insert(0, tr_pr)
    return tr_pr


def set_row_cant_split(tr) -> None:
    """Prevent a single row from splitting across a page boundary."""
    tr_pr = _ensure_tr_pr(tr)
    for old in tr_pr.findall(f"{W}cantSplit"):
        tr_pr.remove(old)
    etree.SubElement(tr_pr, f"{W}cantSplit")


def set_paragraph_keep(p, keep_next: bool) -> None:
    p_pr = p.find(f"{W}pPr")
    if p_pr is None:
        p_pr = etree.Element(f"{W}pPr")
        p.insert(0, p_pr)
    for tag in ("keepNext", "keepLines"):
        for old in p_pr.findall(f"{W}{tag}"):
            p_pr.remove(old)
    etree.SubElement(p_pr, f"{W}keepLines")
    if keep_next:
        etree.SubElement(p_pr, f"{W}keepNext")


def keep_table_together(tbl) -> None:
    """
    Keep a whole table on one page: no row splits, and every row is kept
    with the following row so Word shifts the entire table to the next page
    when it does not fit in the remaining space. Appearance is unchanged.
    """
    rows = tbl.findall(f"{W}tr")
    last_idx = len(rows) - 1
    for row_idx, tr in enumerate(rows):
        set_row_cant_split(tr)
        keep_next = row_idx < last_idx
        for tc in tr.findall(f"{W}tc"):
            for p in tc.findall(f"{W}p"):
                set_paragraph_keep(p, keep_next=keep_next)


def set_paragraph_text(paragraph, text: str, bold: bool | None = None) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
        if bold is not None:
            paragraph.runs[0].bold = bold
    else:
        run = paragraph.add_run(text)
        if bold is not None:
            run.bold = bold


def set_tc_text(tc, text: str, bold: bool | None = None) -> None:
    texts = tc.findall(f".//{W}t")
    value = "" if text is None else str(text)
    if texts:
        texts[0].text = value
        for node in texts[1:]:
            node.text = ""
    else:
        p = tc.find(f"{W}p")
        if p is None:
            p = etree.SubElement(tc, f"{W}p")
        t = etree.SubElement(p, f"{W}t")
        t.text = value
    if bold is not None:
        for r in tc.findall(f".//{W}r"):
            rpr = r.find(f"{W}rPr")
            if rpr is None:
                rpr = etree.SubElement(r, f"{W}rPr")
            b = rpr.find(f"{W}b")
            if bold:
                if b is None:
                    etree.SubElement(rpr, f"{W}b")
            else:
                if b is not None:
                    rpr.remove(b)


def format_amount(value) -> str:
    if value is None or value == "":
        return ""
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    if number == 0:
        return ""
    return f"{int(round(number)):,}/="


def format_visit_date(value: str) -> str:
    if not value:
        return ""
    raw = str(value).strip()
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d/%m/%y"):
        try:
            dt = datetime.strptime(raw[:10], fmt)
            return f"{dt.day:02d}/{dt.month:02d}/{str(dt.year)[-2:]}"
        except ValueError:
            continue
    return raw


def format_export_date(value: str) -> str:
    if not value:
        return datetime.now().strftime("%d/%m/%Y")
    raw = str(value).strip()
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d/%m/%y"):
        try:
            return datetime.strptime(raw[:10], fmt).strftime("%d/%m/%Y")
        except ValueError:
            continue
    return raw


def parse_visit_sort_key(value: str):
    raw = format_visit_date(value)
    cleaned = re.sub(r"/+", "/", raw.strip())
    for fmt in ("%d/%m/%y", "%d/%m/%Y"):
        try:
            return datetime.strptime(cleaned, fmt)
        except ValueError:
            continue
    return datetime.min


def make_table_element(source_tbl, header_tr, data_trs):
    new_tbl = etree.Element(f"{W}tbl")
    for tag in ("tblPr", "tblGrid"):
        el = source_tbl.find(f"{W}{tag}")
        if el is not None:
            new_tbl.append(copy.deepcopy(el))
    new_tbl.append(copy.deepcopy(header_tr))
    for tr in data_trs:
        new_tbl.append(copy.deepcopy(tr))
    return new_tbl


def fill_data_row(template_tr, values, bold_treatment: bool = False, bold_amount: bool = False):
    tr = copy.deepcopy(template_tr)
    tcs = tr.findall(f"{W}tc")
    for idx, val in enumerate(values):
        if idx >= len(tcs):
            break
        bold = None
        if idx == 6 and bold_treatment:
            bold = True
        if idx == 7 and bold_amount:
            bold = True
        set_tc_text(tcs[idx], val, bold=bold)
    return tr


def build_visit_rows(visit, line_tr, discount_tr, total_tr):
    base = [
        format_visit_date(visit.get("visitDate", "")),
        str(visit.get("patientId", "") or ""),
        format_title_case_words(visit.get("patientName", "")),
        format_age(visit.get("age", "")),
        format_sex(visit.get("sex", "")),
        format_title_case_words(visit.get("diagnosis", "")),
    ]
    empty_base = ["", "", "", "", "", ""]
    rows = []
    lines = visit.get("lines") or []
    if not lines:
        lines = [{"label": "Medical bills", "amount": visit.get("visitTotal", 0)}]
    for index, line in enumerate(lines):
        row_base = base if index == 0 else empty_base
        rows.append(
            fill_data_row(
                line_tr,
                row_base
                + [
                    format_title_case_words(line.get("label", "") or ""),
                    format_amount(line.get("amount")),
                ],
            )
        )
    discount = visit.get("discount", 0)
    discount_text = ""
    try:
        if float(discount) > 0:
            discount_text = str(int(round(float(discount))))
    except (TypeError, ValueError):
        discount_text = str(discount or "")
    rows.append(
        fill_data_row(
            discount_tr,
            empty_base + ["Discount", discount_text],
            bold_treatment=True,
        )
    )
    rows.append(
        fill_data_row(
            total_tr,
            empty_base + ["Total Amount:", format_amount(visit.get("visitTotal", 0))],
            bold_treatment=True,
            bold_amount=True,
        )
    )
    return rows


def build_invoice(data: dict, template_path: Path, output_path: Path) -> None:
    template = Document(str(template_path))
    source_table = template.tables[0]
    source_tbl = source_table._tbl
    header_tr = source_tbl.findall(f"{W}tr")[0]
    line_tr = source_tbl.findall(f"{W}tr")[1]
    discount_tr = source_tbl.findall(f"{W}tr")[2]
    total_tr = source_tbl.findall(f"{W}tr")[3]
    grand_total_tbl = template.tables[-1]._tbl
    grand_total_tr = grand_total_tbl.findall(f"{W}tr")[0]

    doc = Document(str(template_path))
    body = doc.element.body

    for tbl in list(doc.tables):
        body.remove(tbl._tbl)

    while len(doc.paragraphs) > 3:
        body.remove(doc.paragraphs[-1]._element)

    facility = format_title_case_words(str(data.get("facilityName") or "VENERANDA MEDICAL"))
    group_name = format_title_case_words(str(data.get("groupName") or ""))
    period = str(data.get("periodLabel") or "").upper()
    export_date = format_export_date(data.get("exportDate"))
    grand_total = format_amount(data.get("grandTotal", 0))

    set_paragraph_text(doc.paragraphs[0], f"INVOICE FOR MEDICAL BILLS {period}", bold=True)
    set_paragraph_text(
        doc.paragraphs[1],
        f"FROM: Hospital/Health Center Name: {facility}                                 DATE: {export_date}",
    )
    set_paragraph_text(
        doc.paragraphs[2],
        f"TO: {group_name}                                                            Total: {grand_total}",
    )

    visits = list(data.get("visits") or [])
    visits.sort(key=lambda item: parse_visit_sort_key(item.get("visitDate", "")), reverse=True)

    insert_pos = list(body).index(doc.paragraphs[2]._element) + 1

    for visit in visits:
        visit_rows = build_visit_rows(visit, line_tr, discount_tr, total_tr)
        new_tbl = make_table_element(source_tbl, header_tr, visit_rows)
        merge_patient_info_columns(new_tbl)
        keep_table_together(new_tbl)
        body.insert(insert_pos, new_tbl)
        insert_pos += 1
        body.insert(insert_pos, etree.Element(f"{W}p"))
        insert_pos += 1

    gt_tbl = etree.Element(f"{W}tbl")
    for tag in ("tblPr", "tblGrid"):
        el = grand_total_tbl.find(f"{W}{tag}")
        if el is not None:
            gt_tbl.append(copy.deepcopy(el))
    gt_row = copy.deepcopy(grand_total_tr)
    gt_cells = gt_row.findall(f"{W}tc")
    if gt_cells:
        set_tc_text(gt_cells[-1], grand_total, bold=True)
    gt_tbl.append(gt_row)
    keep_table_together(gt_tbl)
    body.insert(insert_pos, gt_tbl)

    doc.save(str(output_path))


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--data", required=True, help="Path to JSON payload")
    args = parser.parse_args()

    payload = json.loads(Path(args.data).read_text(encoding="utf-8-sig"))
    build_invoice(payload, Path(args.template), Path(args.output))
    print(f"Saved: {args.output}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
